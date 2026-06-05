#!/usr/bin/env python3
"""
#genai: Script to generate sample test files for DocSeva manual and automated testing.
Run: python generate_samples.py
Requires: python-docx, openpyxl, fpdf2, Pillow
"""
from __future__ import annotations

from pathlib import Path

BASE = Path(__file__).parent

# ── DOCX helper ───────────────────────────────────────────────────────────────

def _heading(doc, text, level=1):
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return p

def _para(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    return p

def _table_row(table, *cells, bold=False):
    row = table.add_row()
    for i, cell_text in enumerate(cells):
        cell = row.cells[i]
        cell.text = str(cell_text)
        if bold:
            for run in cell.paragraphs[0].runs:
                run.bold = True
    return row

# ── 1. sample_quotation.docx ──────────────────────────────────────────────────

def make_quotation():
    from docx import Document
    doc = Document()
    _heading(doc, "QUOTATION", 1)
    _para(doc, "Ref No: QT-2024-001")
    _para(doc, "Date: 15-05-2024")
    _para(doc, "Valid Until: 15-06-2024")
    doc.add_paragraph()
    _para(doc, "To,")
    _para(doc, "ABC Infrastructure Pvt Ltd")
    _para(doc, "45 Industrial Estate, Phase-2")
    _para(doc, "Pune, Maharashtra - 411018")
    doc.add_paragraph()
    _para(doc, "Subject: Quotation for Supply of Survey Equipment", bold=True)
    doc.add_paragraph()
    _para(doc, "Dear Sir/Madam,")
    _para(doc, "We are pleased to submit our quotation for the items listed below:")
    doc.add_paragraph()

    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(["Sr. No.", "Description", "Qty", "Unit Price (₹)", "Total (₹)"]):
        hdr[i].text = h
        for run in hdr[i].paragraphs[0].runs:
            run.bold = True

    items = [
        ("1", "Total Station (5 Second Accuracy)", "2 Units", "85,000", "1,70,000"),
        ("2", "Auto Level Instrument", "3 Units", "22,000", "66,000"),
        ("3", "GPS/GNSS Receiver (Dual Frequency)", "1 Unit", "1,45,000", "1,45,000"),
        ("4", "Tripod Stand (Heavy Duty Aluminium)", "5 Nos", "4,500", "22,500"),
        ("5", "Survey Staff (5 Meter Telescopic)", "4 Nos", "2,800", "11,200"),
    ]
    for row_data in items:
        row = table.add_row()
        for i, val in enumerate(row_data):
            row.cells[i].text = val

    doc.add_paragraph()
    _para(doc, "Subtotal: ₹4,14,700")
    _para(doc, "GST @18%: ₹74,646")
    _para(doc, "Grand Total: ₹4,89,346")
    doc.add_paragraph()
    _para(doc, "Terms & Conditions:", bold=True)
    _para(doc, "• Delivery within 15 working days from order confirmation.")
    _para(doc, "• Payment: 50% advance, balance before dispatch.")
    _para(doc, "• Warranty: 1 year from date of supply.")
    doc.add_paragraph()
    _para(doc, "For Sanmati Enterprises")
    _para(doc, "Authorised Signatory")

    path = BASE / "quotation" / "sample_quotation.docx"
    doc.save(str(path))
    print(f"✅ {path}")


# ── 2. sample_format_template.docx ────────────────────────────────────────────

def make_format_template():
    from docx import Document
    from docx.shared import Pt
    doc = Document()
    _heading(doc, "QUOTATION / OFFER LETTER", 1)
    doc.add_paragraph()
    _para(doc, "Ref: [REF_NO]        Date: [DATE]")
    _para(doc, "Valid Till: [VALID_UNTIL]")
    doc.add_paragraph()
    _para(doc, "To,")
    _para(doc, "[CLIENT NAME]")
    _para(doc, "[CLIENT ADDRESS]")
    doc.add_paragraph()
    _para(doc, "Subject: [SUBJECT]", bold=True)
    doc.add_paragraph()
    _para(doc, "Dear Sir / Madam,")
    _para(doc, "With reference to your enquiry, we hereby submit our competitive offer as under:")
    doc.add_paragraph()

    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    headers = ["S.No.", "Particulars", "Qty", "Rate (₹)", "Amount (₹)"]
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h

    for i in range(1, 4):
        row = table.add_row()
        row.cells[0].text = str(i)
        row.cells[1].text = "[Item Description]"
        row.cells[2].text = "[Qty]"
        row.cells[3].text = "[Rate]"
        row.cells[4].text = "[Amount]"

    doc.add_paragraph()
    _para(doc, "Sub Total  : ₹ [SUBTOTAL]")
    _para(doc, "GST @18%   : ₹ [GST]")
    _para(doc, "Grand Total: ₹ [GRAND_TOTAL]")
    doc.add_paragraph()
    _para(doc, "Note: All prices are exclusive of GST unless mentioned.")
    doc.add_paragraph()
    _para(doc, "Thanking you,")
    _para(doc, "For [COMPANY NAME]")
    _para(doc, "Authorised Signatory")

    path = BASE / "quotation" / "sample_format_template.docx"
    doc.save(str(path))
    print(f"✅ {path}")


# ── 3. sample_invoice_complete.docx ───────────────────────────────────────────

def make_invoice_complete():
    from docx import Document
    doc = Document()
    _heading(doc, "TAX INVOICE", 1)
    doc.add_paragraph()
    _para(doc, "Bill To:")
    _para(doc, "XYZ Construction Ltd")
    _para(doc, "Plot No. 12, Sector 5, Noida, UP - 201301")
    _para(doc, "GSTIN: 09ABCDE1234F1Z5")
    _para(doc, "State: Uttar Pradesh (09)")
    doc.add_paragraph()
    _para(doc, "Ship To:")
    _para(doc, "XYZ Construction Ltd - Site Office")
    _para(doc, "Plot No. 12, Sector 5, Noida, UP - 201301")
    doc.add_paragraph()

    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    for i, h in enumerate(["S.No.", "Item Description (HSN)", "Qty", "Unit Rate", "Amount"]):
        table.rows[0].cells[i].text = h

    items = [
        ("1", "Cement OPC 53 Grade (25 kg Bag) (HSN: 2523)", "200 Bags", "380.00", "76,000.00"),
        ("2", "TMT Steel Bars Fe-500 Grade 10mm (HSN: 7214)", "500 Kg", "68.00", "34,000.00"),
        ("3", "Coarse Aggregate 20mm (HSN: 2517)", "10 CFT", "45.00", "450.00"),
        ("4", "Binding Wire 16 Gauge (HSN: 7217)", "50 Kg", "85.00", "4,250.00"),
    ]
    for row_data in items:
        row = table.add_row()
        for i, val in enumerate(row_data):
            row.cells[i].text = val

    doc.add_paragraph()
    _para(doc, "Subtotal: ₹1,14,700.00")
    _para(doc, "GST @18%: ₹20,646.00")
    _para(doc, "Grand Total: ₹1,35,346.00")

    path = BASE / "invoice" / "sample_invoice_complete.docx"
    doc.save(str(path))
    print(f"✅ {path}")


# ── 4. sample_invoice_no_hsn.docx ─────────────────────────────────────────────

def make_invoice_no_hsn():
    from docx import Document
    doc = Document()
    _heading(doc, "TAX INVOICE / BILL", 1)
    doc.add_paragraph()
    _para(doc, "Bill To:")
    _para(doc, "Rajesh Traders")
    _para(doc, "15 Gandhi Nagar, Jaipur, Rajasthan - 302001")
    _para(doc, "GSTIN: 08FGHIJ5678K2L6")
    doc.add_paragraph()

    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    for i, h in enumerate(["Sr.", "Description", "Qty", "Amount"]):
        table.rows[0].cells[i].text = h

    items = [
        ("1", "Industrial Water Pump 5HP", "3 Nos", "42,000"),
        ("2", "Flexible Hose Pipe 1 inch 10m", "10 Rolls", "3,500"),
        ("3", "Pressure Gauge 0-10 Bar", "6 Nos", "2,400"),
        ("4", "Ball Valve 1 inch SS", "20 Nos", "8,000"),
        ("5", "Cable Tray Perforated 100mm", "30 Mtrs", "9,000"),
    ]
    for row_data in items:
        row = table.add_row()
        for i, val in enumerate(row_data):
            row.cells[i].text = val

    doc.add_paragraph()
    _para(doc, "Note: HSN codes not mentioned — to be confirmed.")
    _para(doc, "Total Amount: ₹64,900")

    path = BASE / "invoice" / "sample_invoice_no_hsn.docx"
    doc.save(str(path))
    print(f"✅ {path}")


# ── 5. sample_invoice_no_billto.docx ──────────────────────────────────────────

def make_invoice_no_billto():
    from docx import Document
    doc = Document()
    _heading(doc, "SUPPLY ORDER / INVOICE", 1)
    doc.add_paragraph()
    _para(doc, "Invoice for supply of IT equipment as per PO reference PO-2024-089")
    doc.add_paragraph()

    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    for i, h in enumerate(["S.No.", "Item (HSN)", "Qty", "Rate", "Total"]):
        table.rows[0].cells[i].text = h

    items = [
        ("1", "Laptop Dell Latitude 5540 i5 (HSN: 8471)", "5 Nos", "72,000", "3,60,000"),
        ("2", "Wireless Mouse Logitech M330 (HSN: 8471)", "10 Nos", "1,200", "12,000"),
        ("3", "USB-C Hub 7-in-1 (HSN: 8504)", "5 Nos", "2,500", "12,500"),
        ("4", "Laptop Bag 15.6 inch (HSN: 4202)", "5 Nos", "1,800", "9,000"),
    ]
    for row_data in items:
        row = table.add_row()
        for i, val in enumerate(row_data):
            row.cells[i].text = val

    doc.add_paragraph()
    _para(doc, "Subtotal: ₹3,93,500.00")
    _para(doc, "GST @18%: ₹70,830.00")
    _para(doc, "Grand Total: ₹4,64,330.00")

    path = BASE / "invoice" / "sample_invoice_no_billto.docx"
    doc.save(str(path))
    print(f"✅ {path}")


# ── 6. sample_invoice_multi_table.docx ────────────────────────────────────────

def make_invoice_multi_table():
    from docx import Document
    doc = Document()
    _heading(doc, "CONSOLIDATED BILL OF SUPPLY", 1)
    _para(doc, "Bill To: Maharashtra State Electricity Board")
    _para(doc, "Vidyut Bhavan, Prakash Nagar, Nagpur - 440001")
    _para(doc, "GSTIN: 27MHSEB0001A1Z9")
    doc.add_paragraph()
    _para(doc, "Section A: Civil Works", bold=True)

    tbl1 = doc.add_table(rows=1, cols=4)
    tbl1.style = "Table Grid"
    for i, h in enumerate(["Sr.", "Description (HSN)", "Qty", "Amount"]):
        tbl1.rows[0].cells[i].text = h
    for row_data in [
        ("1", "Excavation of Trench 1m x 1m (HSN: 9954)", "50 Mtrs", "15,000"),
        ("2", "Backfilling with Murrum (HSN: 9954)", "50 Mtrs", "7,500"),
        ("3", "PCC M10 Bedding (HSN: 9954)", "10 Cu M", "12,000"),
    ]:
        row = tbl1.add_row()
        for i, v in enumerate(row_data):
            row.cells[i].text = v

    doc.add_paragraph()
    _para(doc, "Section B: Electrical Works", bold=True)
    tbl2 = doc.add_table(rows=1, cols=4)
    tbl2.style = "Table Grid"
    for i, h in enumerate(["Sr.", "Description (HSN)", "Qty", "Amount"]):
        tbl2.rows[0].cells[i].text = h
    for row_data in [
        ("1", "XLPE Cable 3.5 Core 95 Sqmm (HSN: 8544)", "200 Mtrs", "1,60,000"),
        ("2", "Cable Jointing Kit HV (HSN: 8544)", "4 Nos", "24,000"),
        ("3", "Earth Pit GI Plate (HSN: 8534)", "8 Nos", "16,000"),
    ]:
        row = tbl2.add_row()
        for i, v in enumerate(row_data):
            row.cells[i].text = v

    doc.add_paragraph()
    _para(doc, "Grand Total (All Sections): ₹2,34,500")
    _para(doc, "GST @18%: ₹42,210")
    _para(doc, "Net Payable: ₹2,76,710")

    path = BASE / "invoice" / "sample_invoice_multi_table.docx"
    doc.save(str(path))
    print(f"✅ {path}")


# ── 7. sample_products.xlsx ───────────────────────────────────────────────────

def make_excel():
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment

    wb = Workbook()
    ws = wb.active
    ws.title = "Product Catalog"

    headers = ["S.No.", "Product Name", "Category", "Unit", "Unit Price (₹)", "GST %", "HSN Code", "Stock"]
    for col, h in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col, value=h)
        cell.font = Font(bold=True)
        cell.fill = PatternFill("solid", fgColor="4472C4")
        cell.font = Font(bold=True, color="FFFFFF")
        cell.alignment = Alignment(horizontal="center")

    products = [
        (1, "Industrial Gear Pump 5HP", "Pumps", "Nos", 42000, 18, "8413", 15),
        (2, "Centrifugal Pump 3HP SS", "Pumps", "Nos", 28000, 18, "8413", 8),
        (3, "Pressure Switch 10 Bar", "Controls", "Nos", 1800, 18, "8536", 50),
        (4, "Flow Meter Digital 2 inch", "Instruments", "Nos", 8500, 18, "9026", 20),
        (5, "Ball Valve SS 1 inch", "Valves", "Nos", 950, 18, "8481", 200),
        (6, "Gate Valve 2 inch CI", "Valves", "Nos", 2400, 18, "8481", 75),
        (7, "Pipe Flange 4 inch PN16", "Fittings", "Nos", 1200, 18, "7307", 100),
        (8, "Gasket Sheet 3mm", "Fittings", "Meters", 450, 18, "8484", 500),
        (9, "Electric Motor 7.5 HP 4 Pole", "Motors", "Nos", 18500, 18, "8501", 12),
        (10, "VFD Drive 5HP 415V", "Drives", "Nos", 14000, 18, "8504", 10),
    ]

    for row_data in products:
        ws.append(row_data)

    for col in ws.columns:
        max_len = max(len(str(cell.value or "")) for cell in col)
        ws.column_dimensions[col[0].column_letter].width = max_len + 4

    path = BASE / "excel" / "sample_products.xlsx"
    wb.save(str(path))
    print(f"✅ {path}")


# ── 8. sample_gst_invoice.docx ────────────────────────────────────────────────

def make_gst_invoice():
    from docx import Document
    doc = Document()
    _heading(doc, "TAX INVOICE", 1)
    doc.add_paragraph()
    _para(doc, "Supplier: Sanmati Enterprises")
    _para(doc, "GSTIN: 27ABCDE1234F1Z5")
    _para(doc, "Address: 15 Industrial Area, Pune, Maharashtra - 411018")
    _para(doc, "Invoice No: INV-2024-089   Date: 20-05-2024")
    doc.add_paragraph()
    _para(doc, "Bill To:")
    _para(doc, "Rajesh Infrastructure Pvt Ltd")
    _para(doc, "45 MG Road, Bengaluru, Karnataka - 560001")
    _para(doc, "GSTIN: 29FGHIJ5678K2L6  State: Karnataka (29)")
    doc.add_paragraph()

    table = doc.add_table(rows=1, cols=6)
    table.style = "Table Grid"
    for i, h in enumerate(["S.No.", "Description", "HSN", "Qty", "Rate", "Amount"]):
        table.rows[0].cells[i].text = h

    items = [
        ("1", "TMT Bars Fe-500 10mm", "7214", "1000 Kg", "68", "68,000"),
        ("2", "Cement OPC 53 Grade", "2523", "500 Bags", "380", "1,90,000"),
    ]
    for row_data in items:
        row = table.add_row()
        for i, val in enumerate(row_data):
            row.cells[i].text = val

    doc.add_paragraph()
    _para(doc, "Taxable Value  : ₹2,58,000")
    _para(doc, "IGST @18%      : ₹46,440")
    _para(doc, "Total Invoice  : ₹3,04,440")
    _para(doc, "")
    _para(doc, "Amount in Words: Rupees Three Lakh Four Thousand Four Hundred Forty Only")
    _para(doc, "Place of Supply: Karnataka (29)")

    path = BASE / "invoice" / "sample_gst_invoice.docx"
    doc.save(str(path))
    print(f"✅ {path}")


# ── 9. sample_product_image.jpg ───────────────────────────────────────────────

def make_product_image():
    from PIL import Image, ImageDraw, ImageFont
    img = Image.new("RGB", (800, 600), color=(240, 248, 255))
    draw = ImageDraw.Draw(img)

    draw.rectangle([20, 20, 780, 580], outline=(70, 130, 180), width=3)
    draw.rectangle([50, 50, 750, 550], fill=(255, 255, 255), outline=(200, 200, 200), width=1)

    draw.rectangle([150, 80, 650, 300], fill=(200, 220, 240), outline=(100, 150, 200), width=2)
    for i in range(160, 641, 30):
        draw.line([(i, 90), (i + 20, 290)], fill=(180, 200, 220), width=1)

    try:
        font_big = ImageFont.truetype("/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf", 32)
        font_med = ImageFont.truetype("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf", 20)
        font_sm = ImageFont.truetype("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf", 16)
    except Exception:
        font_big = font_med = font_sm = None

    draw.text((400, 330), "Industrial Gear Pump 5HP", fill=(30, 30, 80), anchor="mm", font=font_big)
    draw.text((400, 375), "Model: IGP-500 | HSN: 8413", fill=(80, 80, 120), anchor="mm", font=font_med)
    draw.text((400, 415), "Flow Rate: 500 LPH  |  Head: 30m  |  Motor: 5HP 3-Phase", fill=(100, 100, 100), anchor="mm", font=font_sm)
    draw.text((400, 445), "Price: ₹42,000/- + GST @18%", fill=(180, 40, 40), anchor="mm", font=font_med)
    draw.text((400, 510), "Sanmati Enterprises | +91-98765-43210", fill=(120, 120, 120), anchor="mm", font=font_sm)

    path = BASE / "images" / "sample_product.jpg"
    img.save(str(path), "JPEG", quality=95)
    print(f"✅ {path}")


# ── README ────────────────────────────────────────────────────────────────────

def make_readme():
    content = """\
# DocSeva Test Sample Files

Sample input files to test all bot features. Use these files in Telegram to verify each scenario.

## Files at a Glance

| File | Use For |
|------|---------|
| `quotation/sample_quotation.docx` | Sister Quotation conversion — this is the SOURCE file (input) |
| `quotation/sample_format_template.docx` | Sister Quotation — upload this as the FORMAT TEMPLATE |
| `invoice/sample_invoice_complete.docx` | Bill to Make — happy path (all details present) |
| `invoice/sample_invoice_no_hsn.docx` | Bill to Make — triggers HSN prompt (no HSN codes) |
| `invoice/sample_invoice_no_billto.docx` | Bill to Make — triggers BillTo/ShipTo prompt |
| `invoice/sample_invoice_multi_table.docx` | Bill to Make — multi-table stress test |
| `invoice/sample_gst_invoice.docx` | GST Validate feature |
| `excel/sample_products.xlsx` | Convert to DOCX from Excel |
| `images/sample_product.jpg` | Product Catalog PDF generation |

## Quick Start

1. Open Telegram and start the DocSeva bot.
2. Send any file above.
3. Choose the feature from the action buttons.

Refer to `TESTING_GUIDE.md` in the project root for step-by-step instructions.
"""
    path = BASE / "README.md"
    path.write_text(content)
    print(f"✅ {path}")


# ── main ──────────────────────────────────────────────────────────────────────

if __name__ == "__main__":
    print("Generating DocSeva sample test files…\n")
    make_quotation()
    make_format_template()
    make_invoice_complete()
    make_invoice_no_hsn()
    make_invoice_no_billto()
    make_invoice_multi_table()
    make_excel()
    make_gst_invoice()
    make_product_image()
    make_readme()
    print("\n✅ All sample files created in ./test-samples/")
