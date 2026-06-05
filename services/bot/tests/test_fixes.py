"""
#genai: Tests covering the KI-* fixes from KNOWN_ISSUES.md.
"""
from __future__ import annotations

import json
from pathlib import Path
from unittest.mock import MagicMock, patch

import pytest


# ── KI-04: per-item GST rate + split summary ─────────────────────────────────

class TestPerItemGstRate:
    """KI-04: items can carry their own gst_rate; summary splits by rate."""

    def test_normalize_uses_default_rate_when_missing(self):
        from app.processors.bill_to_make import _normalize_bill_data
        data = {
            "items": [
                {"sno": "1", "name": "A", "unit_cost": 100, "amount": 100},
                {"sno": "2", "name": "B", "unit_cost": 200, "amount": 200},
            ],
            "gst_rate": 18,
        }
        out = _normalize_bill_data(data)
        assert all(it["gst_rate"] == 18 for it in out["items"])
        assert out["gst_amount"] == pytest.approx(54.0, rel=1e-3)

    def test_normalize_per_item_rates(self):
        from app.processors.bill_to_make import _normalize_bill_data
        data = {
            "items": [
                {"sno": "1", "name": "A", "unit_cost": 100, "amount": 100, "gst_rate": 5},
                {"sno": "2", "name": "B", "unit_cost": 200, "amount": 200, "gst_rate": 12},
                {"sno": "3", "name": "C", "unit_cost": 300, "amount": 300, "gst_rate": 18},
            ],
            "gst_rate": 18,
        }
        out = _normalize_bill_data(data)
        rates = [b["rate"] for b in out["gst_breakdown"]]
        assert sorted(rates) == [5, 12, 18]
        # 5*1 + 12*2 + 18*3 = 5 + 24 + 54 = 83
        assert out["gst_amount"] == pytest.approx(83.0, rel=1e-3)

    def test_filters_zero_amount_rows(self):
        from app.processors.bill_to_make import _normalize_bill_data
        data = {
            "items": [
                {"sno": "1", "name": "A", "unit_cost": 100, "amount": 100, "gst_rate": 18},
                {"sno": "2", "name": "BLANK", "unit_cost": 0, "amount": 0, "gst_rate": 18},
                {"sno": "3", "name": "B", "unit_cost": 200, "amount": 200, "gst_rate": 18},
            ],
            "gst_rate": 18,
        }
        out = _normalize_bill_data(data)
        assert len(out["items"]) == 2
        assert [it["sno"] for it in out["items"]] == ["1", "2"]

    def test_pdf_generation_with_logo(self, tmp_path, company_profile, parsed_bill_complete):
        from app.processors.bill_to_make import _build_company_info, generate_bill_pdf
        from app.processors.constants import LOGO_PATH
        out = tmp_path / "bill.pdf"
        company_info = _build_company_info(company_profile)
        generate_bill_pdf(
            parsed_bill_complete, "INV-001", "31-05-2024",
            out, company_info,
            logo_path=LOGO_PATH if LOGO_PATH.exists() else None,
        )
        assert out.exists() and out.stat().st_size > 1000

    def test_pdf_multi_rate(self, tmp_path, company_profile):
        from app.processors.bill_to_make import _build_company_info, generate_bill_pdf
        data = {
            "bill_to": {"name": "ABC", "address": "X", "gstin": "NA", "state_code": "07", "state_name": "Delhi"},
            "ship_to": {"name": "ABC", "address": "X", "gstin": "NA", "state_code": "07", "state_name": "Delhi"},
            "items": [
                {"sno": "1", "name": "Service", "hsn": "998719", "unit_cost": 1000, "amount": 1000, "gst_rate": 18},
                {"sno": "2", "name": "Goods",   "hsn": "7304",   "unit_cost": 500,  "amount": 500,  "gst_rate": 12},
            ],
            "gst_rate": 18,
        }
        out = tmp_path / "multi_rate.pdf"
        generate_bill_pdf(data, "INV-MR-1", "01-01-2025", out, _build_company_info(company_profile))
        assert out.exists() and out.stat().st_size > 1000


# ── KI-05: price-adjust guard ────────────────────────────────────────────────

class TestPriceAdjustGuard:
    def test_is_quote_document_true_for_quote(self):
        from app.processors.service import is_quote_document
        from app.processors.models import QuoteDocument, QuoteSection, QuoteItem
        doc = QuoteDocument(
            recipient_name="R", recipient_address_lines=[], subject="S",
            ref_no="", date="", valid_until="",
            sections=[QuoteSection(name="G", items=[QuoteItem("1.", "Item", "1", 100.0, 100.0)])],
        )
        assert is_quote_document(doc) is True

    def test_is_quote_document_false_for_dict(self):
        from app.processors.service import is_quote_document
        assert is_quote_document({"items": []}) is False

    def test_is_quote_document_false_for_none(self):
        from app.processors.service import is_quote_document
        assert is_quote_document(None) is False

    def test_adjust_prices_rejects_non_quote(self):
        from app.processors.service import adjust_prices
        with pytest.raises(ValueError):
            adjust_prices({"not": "a quote"}, 10)

    def test_adjust_prices_increases(self):
        from app.processors.service import adjust_prices
        from app.processors.models import QuoteDocument, QuoteSection, QuoteItem
        doc = QuoteDocument(
            recipient_name="R", recipient_address_lines=[], subject="S",
            ref_no="", date="", valid_until="",
            sections=[QuoteSection(name="G", items=[QuoteItem("1.", "X", "2", 1000.0, 2000.0)])],
        )
        adjusted = adjust_prices(doc, 10)
        new_price = adjusted.sections[0].items[0].unit_price
        assert 1050 <= new_price <= 1150


# ── KI-10: GST validator merge logic ─────────────────────────────────────────

class TestGstChunkMerge:
    def test_merge_combines_items_and_dedupes(self):
        from app.processors.gst_validator import _merge_validation_results
        r1 = {
            "items": [
                {"sno": "1", "description": "A", "hsn": "9023", "taxable_amount": 100,
                 "gst_invoiced": 18, "gst_correct": 18, "total_invoiced": 118,
                 "math_ok": True, "hsn_valid": True, "hsn_matches_product": True, "gst_rate_correct": True},
            ],
            "totals": {}, "overall_valid": True, "error_count": 0,
        }
        r2 = {
            "items": [
                {"sno": "1", "description": "A", "hsn": "9023", "taxable_amount": 100,
                 "gst_invoiced": 18, "gst_correct": 18, "total_invoiced": 118,
                 "math_ok": True, "hsn_valid": True, "hsn_matches_product": True, "gst_rate_correct": True},
                {"sno": "2", "description": "B", "hsn": "9024", "taxable_amount": 200,
                 "gst_invoiced": 36, "gst_correct": 36, "total_invoiced": 236,
                 "math_ok": True, "hsn_valid": True, "hsn_matches_product": True, "gst_rate_correct": True},
            ],
            "totals": {}, "overall_valid": True, "error_count": 0,
        }
        merged = _merge_validation_results([r1, r2])
        assert len(merged["items"]) == 2
        assert merged["totals"]["total_taxable_invoiced"] == 300
        assert merged["overall_valid"] is True

    def test_merge_handles_empty_list(self):
        from app.processors.gst_validator import _merge_validation_results
        out = _merge_validation_results([])
        assert out["items"] == []
        assert out["overall_valid"] is True

    def test_validate_short_text_calls_llm_once(self):
        from app.processors.gst_validator import validate_gst_invoice
        with patch("openai.OpenAI") as oc:
            inst = oc.return_value
            inst.chat.completions.create.return_value = _mock_completion(json.dumps({
                "items": [], "totals": {}, "overall_valid": True, "error_count": 0
            }))
            validate_gst_invoice("k", "m", "x" * 100)
            assert inst.chat.completions.create.call_count == 1

    def test_validate_long_text_chunked(self):
        from app.processors.gst_validator import validate_gst_invoice
        with patch("openai.OpenAI") as oc:
            inst = oc.return_value
            inst.chat.completions.create.return_value = _mock_completion(json.dumps({
                "items": [], "totals": {}, "overall_valid": True, "error_count": 0
            }))
            validate_gst_invoice("k", "m", "x" * 20000)
            assert inst.chat.completions.create.call_count >= 2


def _mock_completion(content: str):
    m = MagicMock()
    m.choices = [MagicMock()]
    m.choices[0].message.content = content
    return m


# ── KI-12: .xls / .xlsx extract_text ─────────────────────────────────────────

class TestExtractTextExcel:
    def test_xlsx_extraction(self, tmp_path):
        import openpyxl
        path = tmp_path / "sample.xlsx"
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "Data"
        ws.append(["Name", "Qty", "Price"])
        ws.append(["Steel", 10, 1500])
        ws.append(["Bolts", 50, 5])
        wb.save(str(path))

        from app.processors.extractors import extract_text
        text = extract_text(str(path))
        assert "Steel" in text
        assert "Bolts" in text
        assert "Sheet: Data" in text

    def test_xls_extraction(self, tmp_path):
        # Build a tiny .xls via xlwt if present, otherwise skip
        try:
            import xlwt
        except ImportError:
            pytest.skip("xlwt not installed in this environment")
        path = tmp_path / "old.xls"
        wb = xlwt.Workbook()
        ws = wb.add_sheet("Sheet1")
        ws.write(0, 0, "Item")
        ws.write(0, 1, "Cost")
        ws.write(1, 0, "Widget")
        ws.write(1, 1, 99)
        wb.save(str(path))
        from app.processors.extractors import extract_text
        text = extract_text(str(path))
        assert "Widget" in text
        assert "99" in text

    def test_unsupported_raises(self, tmp_path):
        path = tmp_path / "x.weird"
        path.write_text("hello")
        from app.processors.extractors import extract_text
        with pytest.raises(ValueError):
            extract_text(str(path))


# ── KI-13: watermark mode (logo / text) ──────────────────────────────────────

class TestWatermarkModes:
    def test_text_watermark_creates_png(self, tmp_path):
        from PIL import Image
        from app.processors.watermark import add_watermark
        src = tmp_path / "src.png"
        Image.new("RGB", (400, 300), "white").save(str(src))
        out = add_watermark(src, tmp_path / "out", mode="text", text="DRAFT")
        assert out.exists()
        assert out.suffix == ".png"
        assert Image.open(out).size == (400, 300)

    def test_logo_watermark_with_custom_logo(self, tmp_path):
        from PIL import Image
        from app.processors.watermark import add_watermark
        src = tmp_path / "src.png"
        Image.new("RGB", (600, 400), "white").save(str(src))
        logo = tmp_path / "logo.png"
        Image.new("RGBA", (200, 200), (255, 0, 0, 255)).save(str(logo))
        out = add_watermark(src, tmp_path / "wm", mode="logo", logo_path=logo)
        assert out.exists()

    def test_logo_missing_falls_back_to_text(self, tmp_path):
        from PIL import Image
        from app.processors.watermark import add_watermark
        src = tmp_path / "src.png"
        Image.new("RGB", (300, 300), "white").save(str(src))
        out = add_watermark(src, tmp_path / "out", mode="logo", logo_path=tmp_path / "missing.png")
        assert out.exists()


# ── KI-16: SessionStore Redis fallback ───────────────────────────────────────

class TestSessionStoreRedisFallback:
    def test_in_memory_when_redis_url_empty(self, monkeypatch):
        monkeypatch.delenv("REDIS_URL", raising=False)
        from app.session_store import SessionStore, BotState
        store = SessionStore()
        s = store.get("u1")
        s.state = BotState.WAITING_BILL_META
        store.save("u1", s)
        assert store.get("u1").state == BotState.WAITING_BILL_META

    def test_reset_preserves_auth(self):
        from app.session_store import SessionStore, BotState
        store = SessionStore()
        s = store.get("u2")
        s.is_registered = True
        s.user_id = "U2"
        s.org_id = "O"
        s.state = BotState.WAITING_BILL_HSN
        store.save("u2", s)
        new = store.reset("u2")
        assert new.is_registered is True
        assert new.user_id == "U2"
        assert new.state == BotState.IDLE

    def test_session_serialisation_roundtrip(self):
        from app.session_store import UserSession, BotState, _session_to_json, _session_from_json
        s = UserSession()
        s.state = BotState.WAITING_BILL_HSN
        s.org_id = "ORG-1"
        s.is_registered = True
        s.original_filename = "foo.pdf"
        payload = _session_to_json(s)
        restored = _session_from_json(payload)
        assert restored.state == BotState.WAITING_BILL_HSN
        assert restored.is_registered is True
        assert restored.original_filename == "foo.pdf"


# ── KI-01 / KI-09: logo asset bundled and used ───────────────────────────────

class TestLogoAsset:
    def test_logo_path_resolves(self):
        from app.processors.constants import LOGO_PATH
        # The repo assets/logo.png exists in test environments.
        assert LOGO_PATH.name == "logo.png"

    def test_constants_have_company_fallbacks(self):
        from app.processors import constants
        assert hasattr(constants, "CO_NAME")
        assert hasattr(constants, "CO_ADDRESS")
        assert hasattr(constants, "CO_PHONE")


# ── KI-02: health check ──────────────────────────────────────────────────────

class TestHealthCheck:
    def test_default_health_ok(self):
        from app.health import check_all
        report = check_all()
        # In the test environment all three Python deps are installed.
        assert report.pdf_export_ok is True

    def test_get_health_caches(self):
        from app.health import get_health
        r1 = get_health()
        r2 = get_health()
        assert r1 is r2


# ── KI-03: docxtpl template fallback ─────────────────────────────────────────

class TestDocxtplTemplate:
    def test_template_no_placeholders_returns_false(self, tmp_path):
        from docx import Document
        from app.processors.service import _template_has_placeholders
        path = tmp_path / "plain.docx"
        d = Document()
        d.add_paragraph("Just a plain template with no placeholders.")
        d.save(str(path))
        assert _template_has_placeholders(path) is False

    def test_template_with_placeholders_detected(self, tmp_path):
        from docx import Document
        from app.processors.service import _template_has_placeholders
        path = tmp_path / "tmpl.docx"
        d = Document()
        d.add_paragraph("Hello {{ recipient_name }}")
        d.save(str(path))
        assert _template_has_placeholders(path) is True


# ── KI-14: comparison styling smoke test ─────────────────────────────────────

class TestComparisonStyling:
    def test_compare_builds_styled_table(self, tmp_path):
        from app.processors.quotation_compare import _build_docx
        data = {
            "suppliers": ["Supplier A", "Supplier B"],
            "items": [
                {"item_name": "Pipe", "unit": "Nos", "prices": {"Supplier A": 100, "Supplier B": 120},
                 "lowest_price": 100, "lowest_from": "Supplier A", "price_diff_pct": 20.0, "notes": "A is cheaper"},
                {"item_name": "Bolt", "unit": "Nos", "prices": {"Supplier A": 50, "Supplier B": 45},
                 "lowest_price": 45, "lowest_from": "Supplier B", "price_diff_pct": 11.1, "notes": "B is cheaper"},
            ],
            "recommendation": "Mix of suppliers per item.",
        }
        out = tmp_path / "cmp.docx"
        result = _build_docx(data, out)
        assert result.exists() and result.suffix == ".docx"
        # Re-open and confirm the header cells exist
        from docx import Document
        doc = Document(str(result))
        # First table is the comparison table
        assert len(doc.tables) >= 1
        header_row = doc.tables[0].rows[0]
        texts = [c.text for c in header_row.cells]
        assert "Item / Description" in texts
        assert "Supplier A" in texts
