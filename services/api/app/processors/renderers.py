#genai: DOCX renderers for quotation formats with dynamic company profile injection.
from __future__ import annotations

import tempfile
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from app.processors.formats import TargetFormat
from app.processors.models import QuoteDocument, QuoteSection, _safe_qty


# ── helpers ───────────────────────────────────────────────────────────────────

def _set_cell_bg(cell, hex_colour: str) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_colour)
    tcPr.append(shd)


def _add_logo_header(doc: Document, logo_path: Path | None, company_name: str) -> None:
    """Add a header showing the company logo (if available) + name.

    #genai: KI fix — never fall back to a bundled "default" logo in production
    output. If the user has not uploaded one we only show the company name so
    the header doesn't reserve awkward empty space for a missing image.
    """
    try:
        section = doc.sections[0]
        header = section.header
        para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        para.clear()
        if logo_path is not None and Path(logo_path).exists():
            run = para.add_run()
            run.add_picture(str(logo_path), width=Inches(0.55))
            name_run = para.add_run(f"    {company_name.upper()}")
        else:
            name_run = para.add_run(company_name.upper())
        name_run.bold = True
        name_run.font.size = Pt(10)
        name_run.font.color.rgb = RGBColor(0, 0, 0)
    except Exception:
        pass


#genai: KI fix — terms render AFTER the items table, with optional user-supplied
#       terms list. If the user provided any (validity/payment/delivery, GST note
#       etc.) we render those; otherwise we fall back to the legacy static list.
_DEFAULT_TERMS = [
    "Delivery within 4–6 weeks from order confirmation",
    "Payment: 50% advance, balance on delivery",
    "This quotation is valid for 30 days",
    "Warranty as per OEM specifications",
]


def _add_terms_block(doc: Document, quote: QuoteDocument) -> None:
    """Render the validity + terms-and-conditions block below the items table."""
    if quote.valid_until:
        doc.add_paragraph(f"\nValid until: {quote.valid_until}")
    doc.add_paragraph("\nTerms & Conditions:")
    terms = quote.terms_list if quote.terms_list else _DEFAULT_TERMS
    for t_item in terms:
        if t_item and str(t_item).strip():
            doc.add_paragraph(str(t_item).strip(), style="List Bullet")


def _add_footer(doc: Document, profile: dict) -> None:
    parts = [profile.get("display_name") or ""]
    if profile.get("phone"):
        parts.append(f"Ph: {profile['phone']}")
    if profile.get("gstin"):
        parts.append(f"GSTIN: {profile['gstin']}")
    if profile.get("address"):
        parts.append(profile["address"])
    footer_text = "  |  ".join(p for p in parts if p)
    if not footer_text:
        return
    try:
        section = doc.sections[0]
        footer = section.footer
        para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        para.clear()
        run = para.add_run(footer_text)
        run.font.size = Pt(7)
        run.font.color.rgb = RGBColor(120, 120, 120)
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    except Exception:
        pass


def _logo_from_profile(profile: dict, tmp_dir: Path) -> Path | None:
    """Download company logo from MinIO if logo_key set, else return None."""
    logo_key = profile.get("logo_key")
    if not logo_key:
        return None
    try:
        #genai: WS-C — API processors hit MinIO via app.core.storage instead of the
        # bot's old app.storage_client. Same MinIO bucket, same key format.
        from app.core.config import settings as api_settings
        from app.core.storage import download_file

        dest = tmp_dir / "company_logo.png"
        data = download_file(api_settings.minio_bucket_assets, logo_key)
        dest.write_bytes(data)
        return dest
    except Exception:
        return None


def _flat_items(quote: QuoteDocument):
    """Return a flat list of all QuoteItems across all sections."""
    items = []
    for section in quote.sections:
        items.extend(section.items)
    return items


def _today() -> str:
    return datetime.today().strftime("%d/%m/%Y")


# ── Public render entry point ─────────────────────────────────────────────────

def render(quote: QuoteDocument, fmt: TargetFormat, output_path: Path, profile: dict) -> Path:
    """
    Render a QuoteDocument to DOCX at output_path.
    profile: dict from CompanyProfileOut — used for company name, footer, logo.
    """
    with tempfile.TemporaryDirectory() as td:
        tmp = Path(td)
        #genai: KI fix — pass through Optional[Path]; no default-logo fallback
        logo = _logo_from_profile(profile, tmp)
        company = profile.get("display_name") or "DocSeva"
        if fmt == TargetFormat.SV_ENTERPRISES:
            _render_detailed(quote, output_path, company, logo, profile)
        elif fmt == TargetFormat.SANMATI:
            _render_department(quote, output_path, company, logo, profile)
        elif fmt == TargetFormat.NR_SURVEY:
            _render_simple(quote, output_path, company, logo, profile)
        else:
            _render_simple(quote, output_path, company, logo, profile)
    return output_path


# ── Format renderers ──────────────────────────────────────────────────────────

def _render_detailed(quote: QuoteDocument, out: Path, company: str, logo: Path | None, profile: dict) -> None:
    """Detailed table-based format — all sections merged into one table."""
    doc = Document()
    for sec in doc.sections:
        sec.top_margin = Inches(0.5)
        sec.bottom_margin = Inches(0.5)
        sec.left_margin = Inches(0.7)
        sec.right_margin = Inches(0.7)

    _add_logo_header(doc, logo, company)

    h = doc.add_heading("QUOTATION", level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    meta.add_run(f"Date: {quote.date or _today()}")

    recipient = quote.recipient_name or "Sir/Madam"
    addr_lines = "\n".join(quote.recipient_address_lines) if quote.recipient_address_lines else ""
    to_text = f"To,\n{recipient}"
    if addr_lines:
        to_text += f"\n{addr_lines}"
    doc.add_paragraph(to_text)

    subject = quote.subject or "Kind consideration"
    ref = quote.ref_no
    if ref:
        doc.add_paragraph(f"Ref: {ref}")
    doc.add_paragraph(f"Sub: {subject}")
    doc.add_paragraph("With reference to above, we hereby submit our competitive quotation:")

    headers = ["S.No", "Description", "Qty", "Unit Price", "Total"]
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    for i, hdr in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = hdr
        cell.paragraphs[0].runs[0].bold = True
        _set_cell_bg(cell, "1F4E79")
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    grand_total = 0.0
    sno = 1
    for section in quote.sections:
        if len(quote.sections) > 1:
            # Section header row
            sec_row = t.add_row().cells
            merged = sec_row[0]
            for ci in range(1, len(headers)):
                merged = merged.merge(sec_row[ci])
            merged.text = section.name
            merged.paragraphs[0].runs[0].bold = True
            _set_cell_bg(merged, "BDD7EE")

        for item in section.items:
            qty_num = _safe_qty(item.qty)
            total = item.total if item.total else round(qty_num * item.unit_price, 2)
            grand_total += total

            row = t.add_row().cells
            row[0].text = str(sno)
            row[1].text = item.description or ""
            row[2].text = item.qty or "1"
            row[3].text = f"₹ {item.unit_price:,.2f}" if item.unit_price else ""
            row[4].text = f"₹ {total:,.2f}"
            sno += 1

    total_row = t.add_row().cells
    total_row[0].merge(total_row[3])
    total_row[0].text = "GRAND TOTAL"
    total_row[0].paragraphs[0].runs[0].bold = True
    total_row[4].text = f"₹ {grand_total:,.2f}"
    total_row[4].paragraphs[0].runs[0].bold = True

    _add_terms_block(doc, quote)

    _add_footer(doc, profile)
    doc.save(str(out))


def _render_department(quote: QuoteDocument, out: Path, company: str, logo: Path | None, profile: dict) -> None:
    """Department-grouped format — one table per section."""
    doc = Document()
    for sec in doc.sections:
        sec.top_margin = Inches(0.5)
        sec.bottom_margin = Inches(0.5)
        sec.left_margin = Inches(0.7)
        sec.right_margin = Inches(0.7)

    _add_logo_header(doc, logo, company)

    h = doc.add_heading("QUOTATION", level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(f"Date: {quote.date or _today()}")
    doc.add_paragraph(f"To: {quote.recipient_name or 'Valued Customer'}")
    if quote.subject:
        doc.add_paragraph(f"Subject: {quote.subject}")

    grand_total = 0.0
    dept_totals: list[tuple[str, float]] = []

    for section in quote.sections:
        doc.add_heading(section.name, level=2)
        t = doc.add_table(rows=1, cols=5)
        t.style = "Table Grid"
        for hdr, cell in zip(["S.No", "Description", "Qty", "Unit Price", "Amount"], t.rows[0].cells):
            cell.text = hdr
            cell.paragraphs[0].runs[0].bold = True
            _set_cell_bg(cell, "C9A84C")

        dept_total = 0.0
        for idx, item in enumerate(section.items, 1):
            qty_num = _safe_qty(item.qty)
            total = item.total if item.total else round(qty_num * item.unit_price, 2)
            dept_total += total

            row = t.add_row().cells
            row[0].text = str(idx)
            row[1].text = item.description or ""
            row[2].text = item.qty or "1"
            row[3].text = f"₹ {item.unit_price:,.2f}" if item.unit_price else ""
            row[4].text = f"₹ {total:,.2f}"

        sub_row = t.add_row().cells
        sub_row[0].merge(sub_row[3])
        sub_row[0].text = f"Sub-Total ({section.name})"
        sub_row[0].paragraphs[0].runs[0].bold = True
        sub_row[4].text = f"₹ {dept_total:,.2f}"
        sub_row[4].paragraphs[0].runs[0].bold = True

        grand_total += dept_total
        dept_totals.append((section.name, dept_total))

    doc.add_paragraph()
    summary = doc.add_table(rows=1, cols=2)
    summary.style = "Table Grid"
    summary.rows[0].cells[0].text = "Section"
    summary.rows[0].cells[1].text = "Amount"
    for cell in summary.rows[0].cells:
        cell.paragraphs[0].runs[0].bold = True
        _set_cell_bg(cell, "1F4E79")
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
    for dept, total in dept_totals:
        row = summary.add_row().cells
        row[0].text = dept
        row[1].text = f"₹ {total:,.2f}"
    total_row = summary.add_row().cells
    total_row[0].text = "GRAND TOTAL"
    total_row[0].paragraphs[0].runs[0].bold = True
    total_row[1].text = f"₹ {grand_total:,.2f}"
    total_row[1].paragraphs[0].runs[0].bold = True

    _add_terms_block(doc, quote)
    _add_footer(doc, profile)
    doc.save(str(out))


def _render_simple(quote: QuoteDocument, out: Path, company: str, logo: Path | None, profile: dict) -> None:
    """Simple single-table format."""
    doc = Document()
    for sec in doc.sections:
        sec.top_margin = Inches(0.5)
        sec.bottom_margin = Inches(0.5)
        sec.left_margin = Inches(0.7)
        sec.right_margin = Inches(0.7)

    _add_logo_header(doc, logo, company)

    h = doc.add_heading(f"QUOTATION — {company.upper()}", level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER

    info = doc.add_paragraph()
    info.add_run(f"Date: {quote.date or _today()}")
    if quote.recipient_name:
        info.add_run(f"   To: {quote.recipient_name}")

    headers = ["#", "Description", "Qty", "Rate", "Total"]
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    for i, hdr in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = hdr
        cell.paragraphs[0].runs[0].bold = True
        _set_cell_bg(cell, "2E74B5")
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    grand = 0.0
    sno = 1
    for section in quote.sections:
        if len(quote.sections) > 1:
            sec_row = t.add_row().cells
            merged = sec_row[0]
            for ci in range(1, len(headers)):
                merged = merged.merge(sec_row[ci])
            merged.text = section.name
            merged.paragraphs[0].runs[0].bold = True
            _set_cell_bg(merged, "DEEAF1")

        for item in section.items:
            qty_num = _safe_qty(item.qty)
            total = item.total if item.total else round(qty_num * item.unit_price, 2)
            grand += total

            row = t.add_row().cells
            row[0].text = str(sno)
            row[1].text = item.description or ""
            row[2].text = item.qty or "1"
            row[3].text = f"₹ {item.unit_price:,.2f}" if item.unit_price else ""
            row[4].text = f"₹ {total:,.2f}"
            sno += 1

    total_row = t.add_row().cells
    total_row[0].merge(total_row[3])
    total_row[0].text = "TOTAL"
    total_row[0].paragraphs[0].runs[0].bold = True
    total_row[4].text = f"₹ {grand:,.2f}"
    total_row[4].paragraphs[0].runs[0].bold = True

    _add_terms_block(doc, quote)
    _add_footer(doc, profile)
    doc.save(str(out))
