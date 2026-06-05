#genai: Excel (xlsx/xls) to DOCX — each sheet becomes a table with solid black borders.
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


def convert_excel_to_docx(excel_path: Path, output_path: Path) -> Path:
    """
    Convert an Excel workbook to DOCX.
    - Each worksheet is rendered as a Word table with solid pure-black borders
      on every cell edge (top, bottom, left, right, insideH, insideV).
    - The first data row is bolded as a header if the sheet has more than one row.
    - Column values that are whole floats are shown without the decimal part.
    - Output filename is the same stem as the input.
    """
    suffix = excel_path.suffix.lower()
    if suffix == ".xlsx":
        sheets = _load_xlsx(excel_path)
    elif suffix == ".xls":
        sheets = _load_xls(excel_path)
    else:
        raise ValueError(f"Unsupported Excel format: {suffix}. Use .xlsx or .xls.")

    doc = Document()
    _set_narrow_margins(doc)

    for idx, (sheet_name, rows) in enumerate(sheets):
        if idx > 0:
            doc.add_page_break()

        heading = doc.add_heading(sheet_name, level=1)
        for run in heading.runs:
            run.font.size = Pt(13)

        if not rows:
            doc.add_paragraph("(empty sheet)")
            continue

        max_cols = max(len(row) for row in rows)
        table = doc.add_table(rows=len(rows), cols=max_cols)
        _apply_black_borders(table)

        treat_first_as_header = len(rows) > 1

        for i, row in enumerate(rows):
            is_header = treat_first_as_header and i == 0
            for j in range(max_cols):
                cell = table.cell(i, j)
                val = row[j] if j < len(row) else None
                cell.text = _cell_str(val)
                if is_header:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            run.bold = True

    doc.save(str(output_path))
    return output_path


# ---------------------------------------------------------------------------
# Excel loaders
# ---------------------------------------------------------------------------

def _load_xlsx(path: Path) -> list[tuple[str, list[list]]]:
    import openpyxl
    wb = openpyxl.load_workbook(str(path), data_only=True)
    sheets: list[tuple[str, list[list]]] = []
    for name in wb.sheetnames:
        ws = wb[name]
        rows: list[list] = []
        for row in ws.iter_rows():
            row_data = [cell.value for cell in row]
            if any(v is not None for v in row_data):
                rows.append(row_data)
        sheets.append((name, rows))
    return sheets


def _load_xls(path: Path) -> list[tuple[str, list[list]]]:
    import xlrd
    wb = xlrd.open_workbook(str(path))
    sheets: list[tuple[str, list[list]]] = []
    for name in wb.sheet_names():
        ws = wb.sheet_by_name(name)
        rows: list[list] = []
        for r in range(ws.nrows):
            row_data = [ws.cell_value(r, c) for c in range(ws.ncols)]
            if any(v != "" and v is not None for v in row_data):
                rows.append(row_data)
        sheets.append((name, rows))
    return sheets


# ---------------------------------------------------------------------------
# Formatting helpers
# ---------------------------------------------------------------------------

def _cell_str(val: object) -> str:
    if val is None:
        return ""
    # Show whole floats without trailing ".0"
    if isinstance(val, float) and val == int(val):
        return str(int(val))
    return str(val)


def _set_narrow_margins(doc: Document) -> None:
    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(1.5)
        section.right_margin = Cm(1.5)


def _apply_black_borders(table) -> None:
    """Solid pure-black borders on every edge of every cell."""
    tbl = table._tbl
    tbl_pr = tbl.find(qn("w:tblPr"))
    if tbl_pr is None:
        tbl_pr = OxmlElement("w:tblPr")
        tbl.insert(0, tbl_pr)

    tbl_borders = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "6")   # 0.75 pt — visually solid
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), "000000")
        tbl_borders.append(border)

    tbl_pr.append(tbl_borders)
