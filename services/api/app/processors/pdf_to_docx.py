#genai: PDF to DOCX conversion preserving table structure and text layout.
from __future__ import annotations

from pathlib import Path
from typing import Any

from docx import Document
from docx.shared import Pt


# Public alias preserved so call-sites can do `from app.processors.pdf_to_docx import pdf_to_docx`.
def pdf_to_docx(pdf_path: Path, output_path: Path) -> Path:
    """Bot-facing alias for convert_pdf_to_docx."""
    return convert_pdf_to_docx(pdf_path, output_path)


def convert_pdf_to_docx(pdf_path: Path, output_path: Path) -> Path:
    """
    Read a PDF and write a DOCX that mirrors the layout:
    - Tables detected by pdfplumber ruling lines are reproduced as Word tables
      with the "Table Grid" style (visible borders).
    - Text outside tables is reproduced as plain paragraphs, top-to-bottom.
    - Multi-page PDFs: page break between each page in the DOCX.
    """
    import pdfplumber  # lazy import — not needed for every bot feature

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    with pdfplumber.open(str(pdf_path)) as pdf:
        for page_num, page in enumerate(pdf.pages):
            if page_num > 0:
                doc.add_page_break()
            _render_page(doc, page)

    doc.save(str(output_path))
    return output_path


# ---------------------------------------------------------------------------
# Internal helpers
# ---------------------------------------------------------------------------

def _render_page(doc: Document, page: Any) -> None:
    tables = page.find_tables()

    if not tables:
        # Simple path — no tables, just dump extracted text as paragraphs.
        text = page.extract_text() or ""
        for line in text.split("\n"):
            stripped = line.strip()
            if stripped:
                doc.add_paragraph(stripped)
        return

    table_bboxes = [t.bbox for t in tables]

    # Collect words that lie outside every table bounding box.
    words = page.extract_words(keep_blank_chars=False, x_tolerance=2, y_tolerance=2)

    def _in_table(word: dict) -> bool:
        x0, y0, x1, y1 = word["x0"], word["top"], word["x1"], word["bottom"]
        for bbox in table_bboxes:
            if (x0 >= bbox[0] - 4 and x1 <= bbox[2] + 4
                    and y0 >= bbox[1] - 4 and y1 <= bbox[3] + 4):
                return True
        return False

    outside_words = [w for w in words if not _in_table(w)]

    # Group outside words into visual lines (similar top-y within 3 pts).
    text_lines: list[tuple[float, str]] = []
    if outside_words:
        cur_y = outside_words[0]["top"]
        cur_group: list[str] = []
        for w in outside_words:
            if abs(w["top"] - cur_y) <= 3:
                cur_group.append(w["text"])
            else:
                if cur_group:
                    text_lines.append((cur_y, " ".join(cur_group)))
                cur_y = w["top"]
                cur_group = [w["text"]]
        if cur_group:
            text_lines.append((cur_y, " ".join(cur_group)))

    # Merge text lines and tables, ordered top-to-bottom by y position.
    elements: list[tuple[float, str, Any]] = [
        (y, "text", txt) for y, txt in text_lines
    ]
    for tbl in tables:
        elements.append((tbl.bbox[1], "table", tbl))
    elements.sort(key=lambda x: x[0])

    for _, kind, content in elements:
        if kind == "text":
            if content.strip():
                doc.add_paragraph(content.strip())
        else:
            extracted = content.extract()
            if extracted:
                _add_docx_table(doc, extracted)


def _add_docx_table(doc: Document, data: list[list]) -> None:
    rows = [row for row in data if any(c is not None for c in row)]
    if not rows:
        return

    max_cols = max(len(row) for row in rows)
    tbl = doc.add_table(rows=len(rows), cols=max_cols)
    tbl.style = "Table Grid"

    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            if j < max_cols:
                tbl.cell(i, j).text = str(val) if val is not None else ""

    doc.add_paragraph()  # visual gap after table
