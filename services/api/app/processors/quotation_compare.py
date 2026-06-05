#genai: Multi-quotation comparison — LLM extraction + colour-coded DOCX table.
from __future__ import annotations

import json
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from app.processors.constants import (
    DARK_NAVY_HEX,
    GOLD_HEX,
    LIGHT_BLUE_HEX,
    LIGHT_GREEN_HEX,
    LIGHT_RED_HEX,
    LIGHT_GRAY_HEX,
    CO_NAME,
    CO_ADDRESS,
    CO_PHONE,
)

# ── LLM prompt ───────────────────────────────────────────────────────────────

_COMPARE_SYSTEM = """\
You are a procurement analyst. You receive text from multiple supplier quotations.
Your job is to:
1. Identify the supplier name from each quotation header.
2. Extract ALL line items with their description and price.
3. Normalise item names so that the same item across different quotations has
   the same "item_name" key (use the most descriptive name you find).
4. Build a comparison table, then add brief analyst notes per item.

Return ONLY a single valid JSON object — no markdown, no extra text:
{
  "suppliers": ["Supplier A", "Supplier B"],
  "items": [
    {
      "item_name": "Normalised item description",
      "unit": "Nos",
      "prices": {
        "Supplier A": 7500.0,
        "Supplier B": 8200.0
      },
      "lowest_price": 7500.0,
      "lowest_from": "Supplier A",
      "price_diff_pct": 9.3,
      "notes": "Supplier A is cheaper by ~9%; verify warranty terms."
    }
  ],
  "recommendation": "Overall Supplier A offers better pricing on most items.",
  "summary": "Comparison across 2 quotations with 5 line items."
}

Rules:
- If a supplier does not quote an item write null for its price.
- price_diff_pct = (max - min) / min * 100, rounded to 1 decimal.
- Keep notes concise (1–2 sentences max).
"""


# ── Public API ────────────────────────────────────────────────────────────────

def _compare_quotations_impl(
    file_paths: list[Path],
    supplier_names: list[str],
    output_path: Path,
    api_key: str,
    model: str,
    company_name: str | None = None,
    logo_path: Path | None = None,
) -> Path:
    """
    Extract items from every quotation via LLM, then build a comparison DOCX.

    Args:
        file_paths:     List of paths to quotation files (doc/docx/pdf).
        supplier_names: Display names to use in the table (one per file).
        output_path:    Where to write the DOCX.
        api_key:        OpenAI API key.
        model:          OpenAI model name.

    Returns:
        Path to the comparison DOCX file.
    """
    texts = _extract_all_texts(file_paths, supplier_names)
    data = _call_llm(api_key, model, texts)
    return _build_docx(data, output_path, company_name=company_name, logo_path=logo_path)


# ── Text extraction ───────────────────────────────────────────────────────────

def _extract_all_texts(paths: list[Path], names: list[str]) -> str:
    from app.processors.extractors import extract_text

    chunks: list[str] = []
    for path, name in zip(paths, names):
        try:
            text = extract_text(str(path))
        except Exception as exc:
            text = f"[Could not extract: {exc}]"
        chunks.append(f"=== QUOTATION FROM: {name} ===\n{text[:4000]}")
    return "\n\n".join(chunks)


# ── LLM call ─────────────────────────────────────────────────────────────────

def _call_llm(api_key: str, model: str, combined_text: str) -> dict:
    from openai import OpenAI

    client = OpenAI(api_key=api_key)
    resp = client.chat.completions.create(
        model=model,
        messages=[
            {"role": "system", "content": _COMPARE_SYSTEM},
            {"role": "user", "content": combined_text},
        ],
        response_format={"type": "json_object"},
        temperature=0,
    )
    return json.loads(resp.choices[0].message.content)


# ── DOCX builder ─────────────────────────────────────────────────────────────

def _build_docx(
    data: dict,
    output_path: Path,
    company_name: str | None = None,
    logo_path: Path | None = None,
) -> Path:
    doc = Document()
    _set_margins(doc)
    _add_logo_header(doc, company_name=company_name or CO_NAME, logo_path=logo_path)
    _add_title(doc)
    _add_table(doc, data)
    _add_recommendation(doc, data)
    _add_footer(doc)

    out = output_path.with_suffix(".docx")
    doc.save(str(out))
    return out


def _set_margins(doc: Document) -> None:
    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.6)
        section.right_margin = Inches(0.6)


def _add_logo_header(
    doc: Document,
    company_name: str | None = None,
    logo_path: Path | None = None,
) -> None:
    """Add logo + company-name header.

    #genai: KI fix — no longer falls back to the bundled default logo. If the
    user has not uploaded one, only the company name is rendered.
    """
    name = company_name or CO_NAME
    try:
        section = doc.sections[0]
        header = section.header
        para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        para.clear()
        if logo_path is not None and Path(logo_path).exists():
            run = para.add_run()
            run.add_picture(str(logo_path), width=Inches(0.55))
            name_run = para.add_run(f"    {name}")
        else:
            name_run = para.add_run(name)
        name_run.bold = True
        name_run.font.size = Pt(11)
    except Exception:
        pass


def _add_footer(doc: Document) -> None:
    try:
        section = doc.sections[0]
        footer = section.footer
        para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        para.clear()
        run = para.add_run(f"{CO_NAME} | Ph: {CO_PHONE} | {CO_ADDRESS}")
        run.font.size = Pt(7)
        run.font.color.rgb = RGBColor(120, 120, 120)
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    except Exception:
        pass


def _add_title(doc: Document) -> None:
    h = doc.add_heading("Quotation Comparison Report", level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in h.runs:
        run.font.color.rgb = RGBColor.from_string(DARK_NAVY_HEX)
    doc.add_paragraph()


def _add_table(doc: Document, data: dict) -> None:
    suppliers: list[str] = data.get("suppliers", [])
    items: list[dict] = data.get("items", [])
    if not items:
        doc.add_paragraph("No comparable items found.")
        return

    n_sup = len(suppliers)
    # Columns: Item | Sup1 | Sup2 | ... | Lowest Price (winner) | Notes
    n_cols = 1 + n_sup + 1 + 1

    table = doc.add_table(rows=1, cols=n_cols)
    table.style = "Table Grid"

    # Header row
    hdr_cells = table.rows[0].cells
    headers = ["Item / Description"] + suppliers + ["Lowest Price", "Review Notes"]
    for i, h_text in enumerate(headers):
        cell = hdr_cells[i]
        cell.text = h_text
        _set_cell_bg(cell, DARK_NAVY_HEX)
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(255, 255, 255)

    # Data rows
    for row_idx, item in enumerate(items):
        row_cells = table.add_row().cells
        prices = item.get("prices", {})
        lowest_from = item.get("lowest_from", "")
        lowest_price = item.get("lowest_price")

        # Find max price for highlighting
        numeric_prices = [v for v in prices.values() if v is not None]
        max_price = max(numeric_prices) if numeric_prices else None

        # Item name
        row_cells[0].text = str(item.get("item_name", ""))
        _style_cell(row_cells[0], bold=True, size=9)
        bg = LIGHT_GRAY_HEX if row_idx % 2 else "FFFFFF"
        _set_cell_bg(row_cells[0], bg)

        # Supplier prices
        for j, sup in enumerate(suppliers):
            price = prices.get(sup)
            cell = row_cells[1 + j]
            if price is None:
                cell.text = "—"
                _set_cell_bg(cell, bg)
            else:
                cell.text = f"Rs.{price:,.0f}"
                if price == lowest_price and n_sup > 1:
                    _set_cell_bg(cell, LIGHT_GREEN_HEX)
                elif max_price is not None and price == max_price and n_sup > 1:
                    _set_cell_bg(cell, LIGHT_RED_HEX)
                else:
                    _set_cell_bg(cell, bg)
            _style_cell(cell, size=9)

        # Lowest price / winner
        winner_cell = row_cells[1 + n_sup]
        if lowest_price is not None:
            winner_cell.text = f"Rs.{lowest_price:,.0f}\n({lowest_from})"
            _set_cell_bg(winner_cell, LIGHT_GREEN_HEX)
        else:
            winner_cell.text = "—"
        _style_cell(winner_cell, size=9)

        # Notes
        notes_cell = row_cells[1 + n_sup + 1]
        notes_cell.text = str(item.get("notes", ""))
        _set_cell_bg(notes_cell, LIGHT_BLUE_HEX)
        _style_cell(notes_cell, italic=True, size=8)

    # Set column widths
    _set_col_widths(table, n_sup)


def _set_col_widths(table, n_sup: int) -> None:
    """Distribute column widths within A4 usable width (~16 cm)."""
    from docx.oxml.ns import qn
    from docx.shared import Cm

    item_w = Cm(4.5)
    notes_w = Cm(3.5)
    winner_w = Cm(3.0)
    sup_w = Cm(max(1.8, (16.0 - 4.5 - 3.5 - 3.0) / max(n_sup, 1)))

    widths = [item_w] + [sup_w] * n_sup + [winner_w, notes_w]
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcW = OxmlElement("w:tcW")
            tcW.set(qn("w:w"), str(int(widths[i].pt * 20)))
            tcW.set(qn("w:type"), "dxa")
            tcPr.append(tcW)


def _add_recommendation(doc: Document, data: dict) -> None:
    rec = data.get("recommendation") or data.get("summary")
    if not rec:
        return
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run("AI Recommendation: ")
    run.bold = True
    run.font.color.rgb = RGBColor.from_string(DARK_NAVY_HEX)
    p.add_run(str(rec))


# ── Styling helpers ───────────────────────────────────────────────────────────

def _set_cell_bg(cell, hex_color: str) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color.upper())
    tcPr.append(shd)


def _style_cell(
    cell,
    bold: bool = False,
    italic: bool = False,
    size: int = 9,
) -> None:
    for para in cell.paragraphs:
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in para.runs:
            run.bold = bold
            run.italic = italic
            run.font.size = Pt(size)


def compare_quotations(
    files: list[tuple],
    output_path: "Path",
    company_profile: dict | None = None,
    logo_path: Path | None = None,
) -> "Path":
    """
    Bot-facing entry point.
    files: list of (path, display_name) tuples.
    company_profile: optional dict for company name/logo overrides.
    logo_path: optional local path to the user's logo (no default fallback).
    """
    #genai: WS-C — moved into API; use app.core.config.
    from app.core.config import settings as api_settings
    paths = [f[0] for f in files]
    names = [f[1] for f in files]
    company_name = (company_profile or {}).get("display_name") if company_profile else None
    return _compare_quotations_impl(
        file_paths=paths,
        supplier_names=names,
        output_path=output_path,
        api_key=api_settings.openai_api_key,
        model=api_settings.openai_model,
        company_name=company_name,
        logo_path=logo_path,
    )
